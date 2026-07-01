"""文档解析（「文档学习」平行链路）：PDF / 文本 / Markdown / Word → 纯文本段。

范围（对齐 CC-document-learning）：仅 PDF + 纯文本/markdown + Word（.docx）；
图片OCR / 网页 / 音视频不做。抽取为切片器输入段 [{text, location}]：
- pdf：pypdf 逐页（location 带「第 N 页」，供来源溯源）；
- docx：python-docx 逐段落合并（location 空，无稳定页码）；
- md/txt 及其它：utf-8 解码整篇（失败用替换字符兜底，绝不抛异常中断）。

无任何外部服务/密钥依赖，纯本地库（pypdf / python-docx），mock 环境照常可跑。
"""
from __future__ import annotations

import io
import logging

logger = logging.getLogger("app.services.document_parse")

# 支持的扩展名 → 归一化文件类型
_PDF_EXTS = {"pdf"}
_DOCX_EXTS = {"docx"}
_TEXT_EXTS = {"md", "markdown", "txt", "text"}
SUPPORTED_EXTS = _PDF_EXTS | _DOCX_EXTS | _TEXT_EXTS


class UnsupportedDocument(Exception):
    """不支持的文档格式（→ code 1001 / 400）。"""


def file_type_of(filename: str) -> str:
    """由文件名推断归一化类型：pdf|docx|md|txt；不支持则抛 UnsupportedDocument。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in _PDF_EXTS:
        return "pdf"
    if ext in _DOCX_EXTS:
        return "docx"
    if ext in _TEXT_EXTS:
        return "md" if ext in {"md", "markdown"} else "txt"
    raise UnsupportedDocument(ext or filename)


def derive_title(filename: str) -> str:
    """去路径去扩展名，取文件主名作为文档标题。"""
    base = filename.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    return (base.rsplit(".", 1)[0] if "." in base else base) or "未命名文档"


def extract_segments(filename: str, raw: bytes) -> list[dict]:
    """抽取为切片器输入段：[{text, location}]。不支持格式抛 UnsupportedDocument。"""
    ftype = file_type_of(filename)
    if ftype == "pdf":
        return _extract_pdf(raw)
    if ftype == "docx":
        return _extract_docx(raw)
    # md / txt
    text = raw.decode("utf-8", errors="replace")
    return [{"text": text, "location": ""}] if text.strip() else []


def extract_text(filename: str, raw: bytes) -> str:
    """抽取为整篇纯文本（段落间以空行分隔，供 content 落库 / 摘要复用）。"""
    segs = extract_segments(filename, raw)
    return "\n\n".join(s["text"].strip() for s in segs if s.get("text", "").strip())


def _extract_pdf(raw: bytes) -> list[dict]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    segs: list[dict] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            segs.append({"text": text, "location": f"第 {i} 页"})
    return segs


def _extract_docx(raw: bytes) -> list[dict]:
    try:
        import docx  # python-docx
    except Exception as exc:  # noqa: BLE001 未安装 → 明确报不支持，不静默产出空文档
        raise UnsupportedDocument("docx（python-docx 未安装）") from exc

    document = docx.Document(io.BytesIO(raw))
    paras = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(paras)
    return [{"text": text, "location": ""}] if text.strip() else []
